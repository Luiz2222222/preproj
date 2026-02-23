"""
Geração do Relatório de Avaliação do TCC em DOCX.
Abordagem: copia o template oficial e substitui apenas os dados (mala direta).
Preserva toda a formatação original: cores, bordas, fontes, header, assinatura.
"""

import io
import os
import re
from docx import Document

# Caminho do template
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'relatorio_template.docx')

# Meses em português
MESES_PT = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro',
}

# Seções do parecer estruturado
SECOES_PARECER = [
    'Resumo',
    'Introdução/Relevância do Trabalho',
    'Revisão Bibliográfica',
    'Desenvolvimento',
    'Conclusões',
    'Considerações Adicionais',
]


def _fmt(valor, casas=1):
    """Formata número decimal para string com vírgula."""
    if valor is None:
        return '-'
    return f'{float(valor):.{casas}f}'.replace('.', ',')


def _obter_label_curso(curso_key):
    """Retorna o label amigável do curso."""
    mapa = {
        'ENGENHARIA_ELETRICA': 'Engenharia Elétrica',
        'ENGENHARIA_CONTROLE_AUTOMACAO': 'Engenharia de Controle e Automação',
    }
    return mapa.get(curso_key, curso_key or 'N/A')


def _data_por_extenso(data):
    """Formata data como '30 de janeiro de 2026'."""
    return f'{data.day} de {MESES_PT[data.month]} de {data.year}'


def _replace_cell_value(cell, new_text):
    """
    Substitui o valor de uma célula preservando toda a formatação.
    Encontra o run que contém texto e substitui; limpa runs vazios extras.
    """
    for p in cell.paragraphs:
        found = False
        for run in p.runs:
            if run.text.strip() and not found:
                run.text = new_text
                found = True
            elif found:
                run.text = ''
        if found:
            return
    # Se nenhum run tinha texto, colocar no primeiro run disponível
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
    elif cell.paragraphs:
        cell.paragraphs[0].text = new_text


def _extract_section(parecer, secao):
    """Extrai texto de uma seção do parecer estruturado (formato ===Seção===)."""
    if not parecer:
        return ''
    pattern = rf'===\s*{re.escape(secao)}\s*===\s*(.*?)(?====|$)'
    match = re.search(pattern, parecer, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ''


def gerar_relatorio_avaliacao_docx(tcc):
    """
    Gera o Relatório de Avaliação do TCC em DOCX.
    Copia o template oficial e substitui os dados (mala direta).
    Retorna bytes do DOCX gerado.
    """
    from .models import AvaliacaoFase1, AvaliacaoFase2, AgendamentoDefesa
    from users.models import Usuario

    # ===== BUSCAR DADOS =====
    avaliacoes_f1 = list(
        AvaliacaoFase1.objects.filter(tcc=tcc, status__in=['BLOQUEADO', 'CONCLUIDO', 'PENDENTE'])
        .select_related('avaliador')
        .order_by('id')
    )
    # Filtrar apenas avaliações que têm notas preenchidas
    avaliacoes_f1 = [a for a in avaliacoes_f1 if a.calcular_nota_total() is not None]
    avaliacoes_f2 = list(
        AvaliacaoFase2.objects.filter(tcc=tcc, status__in=['BLOQUEADO', 'CONCLUIDO'])
        .select_related('avaliador')
        .order_by('id')
    )

    # Data da defesa
    try:
        agendamento = AgendamentoDefesa.objects.get(tcc=tcc)
        data_defesa = agendamento.data
    except AgendamentoDefesa.DoesNotExist:
        from django.utils import timezone
        data_defesa = timezone.localdate()

    data_defesa_fmt = data_defesa.strftime('%d/%m/%Y')
    data_extenso = _data_por_extenso(data_defesa)

    # Calcular notas
    notas_f1 = [float(a.calcular_nota_total()) for a in avaliacoes_f1] if avaliacoes_f1 else []
    nf1 = sum(notas_f1) / len(notas_f1) if notas_f1 else (float(tcc.nf1) if tcc.nf1 else 0)

    notas_f2 = [float(a.calcular_nota_total()) for a in avaliacoes_f2] if avaliacoes_f2 else []
    nf2_media = sum(notas_f2) / len(notas_f2) if notas_f2 else 0

    nf1_ponderado = 0.6 * nf1
    nf2_ponderado = 0.4 * nf2_media
    nota_final = nf1_ponderado + nf2_ponderado

    resultado = tcc.resultado_final or ('APROVADO' if nota_final >= 6 else 'REPROVADO')

    # Dados do aluno
    aluno = tcc.aluno
    curso_label = _obter_label_curso(aluno.curso)

    # ===== ABRIR TEMPLATE (cópia em memória) =====
    doc = Document(TEMPLATE_PATH)

    # ===== 1. TABELA 0 — Informações =====
    t0 = doc.tables[0]
    _replace_cell_value(t0.rows[0].cells[1], tcc.titulo.upper())
    _replace_cell_value(t0.rows[1].cells[1], curso_label)
    _replace_cell_value(t0.rows[2].cells[1], data_defesa_fmt)
    _replace_cell_value(t0.rows[3].cells[1], aluno.nome_completo)

    # ===== 2. TABELA 1 — Notas Fase I =====
    t1 = doc.tables[1]
    for idx, av in enumerate(avaliacoes_f1[:2]):
        row = t1.rows[2 + idx]
        nota_total = av.calcular_nota_total()
        notas = [
            _fmt(av.nota_resumo), _fmt(av.nota_introducao),
            _fmt(av.nota_revisao), _fmt(av.nota_desenvolvimento),
            _fmt(av.nota_conclusoes), _fmt(nota_total),
        ]
        for col, nota in enumerate(notas, start=1):
            _replace_cell_value(row.cells[col], nota)

    # NF1
    _replace_cell_value(t1.rows[4].cells[6], _fmt(nf1))

    # ===== 3. TABELA 2 — Notas Fase II =====
    t2 = doc.tables[2]
    for idx, av in enumerate(avaliacoes_f2[:3]):
        row = t2.rows[2 + idx]
        nota_total = av.calcular_nota_total()
        notas = [
            _fmt(av.nota_coerencia_conteudo), _fmt(av.nota_qualidade_apresentacao),
            _fmt(av.nota_dominio_tema), _fmt(av.nota_clareza_fluencia),
            _fmt(av.nota_observancia_tempo), _fmt(nota_total),
        ]
        for col, nota in enumerate(notas, start=1):
            _replace_cell_value(row.cells[col], nota)

    # NF2
    _replace_cell_value(t2.rows[5].cells[6], _fmt(nf2_media))

    # ===== 4. TABELA 3 — Avaliadores Fase II =====
    t3 = doc.tables[3]
    for idx, av in enumerate(avaliacoes_f2[:3]):
        avaliador = av.avaliador
        tratamento = avaliador.tratamento or ''
        if tratamento == 'Outro' and avaliador.tratamento_customizado:
            tratamento = avaliador.tratamento_customizado
        nome = f'{tratamento} {avaliador.nome_completo}'.strip()
        _replace_cell_value(t3.rows[idx].cells[1], nome)

    # ===== 5. TABELA 4 — Apuração =====
    t4 = doc.tables[4]
    _replace_cell_value(t4.rows[1].cells[0], _fmt(nf1_ponderado))
    _replace_cell_value(t4.rows[1].cells[1], _fmt(nf2_ponderado))
    _replace_cell_value(t4.rows[1].cells[2], _fmt(nota_final))

    # ===== 6. PARÁGRAFO 5 — Corpo do texto (nota e resultado) =====
    p_corpo = doc.paragraphs[5]
    # Substituir a nota final (run que contém valor numérico como "8,3")
    for run in p_corpo.runs:
        text = run.text.strip()
        # Encontrar o run com a nota (formato N,N)
        if re.match(r'^\d+,\d+$', text):
            run.text = _fmt(nota_final)
        # Encontrar o run com o resultado
        elif text in ('APROVADO', 'REPROVADO'):
            run.text = resultado

    # ===== 7. PARÁGRAFO 7 — Data =====
    # Template: run "Recife, " + runs vazios + run "30 de janeiro de 2026" + run "."
    # Substituir apenas o run que contém a data (formato "DD de mês de AAAA")
    p_data = doc.paragraphs[7]
    for run in p_data.runs:
        if re.match(r'\d+ de \w+ de \d{4}', run.text.strip()):
            run.text = data_extenso
            break

    # ===== 8. APÊNDICE A — Comentários Fase I =====
    # Posições fixas dos parágrafos de comentário para cada avaliador
    # Avaliador #1: comentários em [22, 25, 28, 31, 34, 37]
    # Avaliador #2: comentários em [42, 45, 48, 51, 54, 57] (ou vazios pois podem não ter)
    comentarios_indices = [
        [22, 25, 28, 31, 34, 37],  # Avaliador #1
        [42, 45, 48, 51, 54, 57],  # Avaliador #2
    ]

    for av_idx, av in enumerate(avaliacoes_f1[:2]):
        if av_idx >= len(comentarios_indices):
            break

        indices = comentarios_indices[av_idx]
        parecer = av.parecer or ''

        for sec_idx, secao in enumerate(SECOES_PARECER):
            if sec_idx >= len(indices):
                break

            p_idx = indices[sec_idx]
            if p_idx >= len(doc.paragraphs):
                break

            # Extrair comentário da seção
            comentario = _extract_section(parecer, secao)

            # Se não tem formato estruturado, colocar tudo em "Considerações Adicionais"
            if not comentario and sec_idx == 5 and not any(
                _extract_section(parecer, s) for s in SECOES_PARECER[:5]
            ):
                comentario = parecer.strip()

            # Substituir o parágrafo de comentário
            p = doc.paragraphs[p_idx]
            if p.runs:
                p.runs[0].text = comentario
                for run in p.runs[1:]:
                    run.text = ''
            elif comentario:
                p.text = comentario

    # ===== GERAR BYTES =====
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
